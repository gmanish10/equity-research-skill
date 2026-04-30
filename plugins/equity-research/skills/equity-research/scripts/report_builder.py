"""
Build .docx reports for equity research and portfolio reviews.

STUB — the skill will typically build docx files via the `docx` skill's
document-skeleton approach rather than this helper, but this provides a
minimal path using python-docx directly if preferred.

Install: pip install python-docx

Usage:
    python scripts/report_builder.py --template equity-research --input report.json --output AAPL-research.docx
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
except ImportError:
    Document = None  # type: ignore


DISCLAIMER = (
    "This document is for educational and informational purposes only. "
    "It is not personalized financial advice. "
    "Consult a qualified financial advisor before making investment decisions. "
    "Data delayed ~15 minutes and sourced from Yahoo Finance."
)


def build_equity_research(data: dict, output_path: Path) -> None:
    if Document is None:
        raise RuntimeError("python-docx not installed. pip install python-docx")

    doc = Document()

    # Cover
    doc.add_heading(f"{data.get('company', '?')} ({data.get('ticker', '?')})", level=0)
    doc.add_paragraph(f"{data.get('report_date', '')} | Aggressive-Growth Analysis").italic = True

    # Disclaimer (cover)
    p = doc.add_paragraph()
    run = p.add_run("DISCLAIMER: ")
    run.bold = True
    p.add_run(DISCLAIMER)

    doc.add_page_break()

    # Executive summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(data.get("executive_summary", ""))

    # Sections
    for section_title, key in [
        ("Investment Thesis", "thesis"),
        ("Bull / Base / Bear Cases", "cases"),
        ("Business Model", "business_model"),
        ("Sector Context", "sector"),
        ("Fundamentals", "fundamentals"),
        ("Technicals", "technicals"),
        ("Options", "options"),
        ("News, Ownership, Sentiment", "sentiment"),
        ("Multi-Horizon Ratings", "ratings"),
        ("Risk Register", "risks"),
        ("Upside Catalysts", "catalysts"),
    ]:
        doc.add_heading(section_title, level=1)
        content = data.get(key, "")
        if isinstance(content, str):
            doc.add_paragraph(content)
        elif isinstance(content, list):
            for item in content:
                doc.add_paragraph(str(item), style="List Bullet")

    # Back page disclaimer
    doc.add_page_break()
    doc.add_heading("Disclaimer", level=2)
    doc.add_paragraph(DISCLAIMER)

    doc.save(output_path)


def build_portfolio_review(data: dict, output_path: Path) -> None:
    """Portfolio review stub — mirror equity-research with portfolio sections."""
    # TODO: fill in full portfolio-review template as per references/report-templates/portfolio-review.md
    if Document is None:
        raise RuntimeError("python-docx not installed. pip install python-docx")
    doc = Document()
    doc.add_heading("Portfolio Review", level=0)
    doc.add_paragraph(data.get("executive_summary", ""))
    doc.add_paragraph(DISCLAIMER)
    doc.save(output_path)


def _cli() -> int:
    parser = argparse.ArgumentParser(description="Report builder")
    parser.add_argument(
        "--template",
        choices=["equity-research", "portfolio-review"],
        required=True,
    )
    parser.add_argument("--input", required=True, help="Path to JSON input")
    parser.add_argument("--output", required=True, help="Output .docx path")
    args = parser.parse_args()

    with open(args.input) as f:
        data = json.load(f)

    out = Path(args.output)
    if args.template == "equity-research":
        build_equity_research(data, out)
    else:
        build_portfolio_review(data, out)

    print(f"Wrote {out}")
    return 0


if __name__ == "__main__":
    sys.exit(_cli())
